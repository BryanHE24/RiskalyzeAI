# File path: scripts/create_sample_data.py
import os
from faker import Faker
from docx import Document
from fpdf import FPDF
import random
from pathlib import Path
import logging

fake = Faker()
categories = ["Network Security", "Phishing", "Malware", "Access Control", "Data Leak"]

def generate_ticket(i):
    return {
        "title": f"{random.choice(categories)} Case - {fake.uuid4()[:8]}",
        "desc": f"Reported by {fake.name()}\n\n{fake.paragraph()}\n\nPriority: {random.choice(['Low','Medium','High'])}"
    }

def create_samples(output_dir="database/sample_data", count=100):
    output_dir = Path(output_dir)
    output_dir.mkdir(exist_ok=True, parents=True)
    
    # Clear existing samples
    for f in output_dir.glob("ticket_*"):
        f.unlink()

    logging.info(f"Generating {count*3} sample files...")
    
    for i in range(1, count+1):
        ticket = generate_ticket(i)
        base_name = f"ticket_{i}"
        
        # TXT
        (output_dir / f"{base_name}.txt").write_text(f"{ticket['title']}\n\n{ticket['desc']}")
        
        # PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, f"{ticket['title']}\n\n{ticket['desc']}")
        pdf.output(output_dir / f"{base_name}.pdf")
        
        # DOCX
        doc = Document()
        doc.add_heading(ticket['title'], level=1)
        doc.add_paragraph(ticket['desc'])
        doc.save(output_dir / f"{base_name}.docx")

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    create_samples(count=35)  # Creates 105 files
